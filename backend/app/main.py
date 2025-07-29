"""
FastAPI backend uygulaması
"""
from fastapi import FastAPI, HTTPException, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response
from typing import Dict, Any
from sqlalchemy.orm import Session
import json
import time
import logging

# Logger ayarla
logger = logging.getLogger(__name__)

# Local imports
from .database import engine, get_db, Base
from .models import Contract, Role, RoleQuestionConfig, QuestionType, Question, QuestionConfig, ContractData, SystemInfo, GenerationLog

# Create tables
Base.metadata.create_all(bind=engine)

# Zorluk seviyesi helper fonksiyonları
def get_difficulty_level_by_multiplier(salary_multiplier: float):
    """Maaş katsayısına göre zorluk seviyesi belirle - Profesyonel Rubrik Modeli"""
    if salary_multiplier <= 2:
        return {
            "level": "2x",
            "name": "🟢 ORTA SEVİYE UYGULAYICI",
            "description": "2-4 yıl tecrübe - Günlük operasyonu eksiksiz yürütme, temel optimizasyon",
            "experience_years": "2-4 yıl",
            "focus": "K1: 30% • K2: 40% • K3: 25% • K4: 5% • K5: 0%",
            "question_distribution": {
                "K1_temel_bilgi": 30,
                "K2_uygulamali": 40, 
                "K3_hatacozumleme": 25,
                "K4_tasarim": 5,
                "K5_stratejik_liderlik": 0
            }
        }
    elif salary_multiplier <= 3:
        return {
            "level": "3x",
            "name": "🟡 KIDEMLI UZMAN", 
            "description": "5-8 yıl tecrübe - Çapraz disiplinde hâkimiyet, mentorluk, kritik problem çözümü",
            "experience_years": "5-8 yıl",
            "focus": "K1: 15% • K2: 25% • K3: 35% • K4: 20% • K5: 5%",
            "question_distribution": {
                "K1_temel_bilgi": 15,
                "K2_uygulamali": 25,
                "K3_hatacozumleme": 35,
                "K4_tasarim": 20,
                "K5_stratejik_liderlik": 5
            }
        }
    elif salary_multiplier <= 4:
        return {
            "level": "4x",
            "name": "🟠 MİMAR/TEKNİK LİDER", 
            "description": "≥10 yıl tecrübe - Strateji, büyük ölçekli mimari, metodoloji, ekip & süreç yönetimi",
            "experience_years": "≥10 yıl", 
            "focus": "K1: 5% • K2: 15% • K3: 25% • K4: 35% • K5: 20%",
            "question_distribution": {
                "K1_temel_bilgi": 5,
                "K2_uygulamali": 15,
                "K3_hatacozumleme": 25,
                "K4_tasarim": 35,
                "K5_stratejik_liderlik": 20
            }
        }
    else:
        return {
            "level": "5x",
            "name": "🔴 ENTERPRISE UZMAN",
            "description": "15+ yıl tecrübe - Enterprise mimari, strategik kararlar, teknoloji liderliği", 
            "experience_years": "15+ yıl",
            "focus": "K1: 0% • K2: 10% • K3: 20% • K4: 40% • K5: 30%",
            "question_distribution": {
                "K1_temel_bilgi": 0,
                "K2_uygulamali": 10,
                "K3_hatacozumleme": 20,
                "K4_tasarim": 40,
                "K5_stratejik_liderlik": 30
            }
        }

# Default soru tiplerini ekle
def create_default_question_types(db: Session):
    """Varsayılan soru tiplerini oluştur"""
    default_types = [
        {
            "name": "Mesleki Tecrübe Soruları",
            "description": "Adayın deneyim ve projelerini değerlendiren sorular",
            "code": "professional_experience",
            "order_index": 1
        },
        {
            "name": "Teorik Bilgi Soruları", 
            "description": "Teknik kavramlar ve teorik bilgiyi ölçen sorular",
            "code": "theoretical_knowledge",
            "order_index": 2
        },
        {
            "name": "Pratik Uygulama Soruları",
            "description": "Problem çözme ve uygulama becerilerini test eden sorular", 
            "code": "practical_application",
            "order_index": 3
        }
    ]
    
    for type_data in default_types:
        existing = db.query(QuestionType).filter(QuestionType.code == type_data["code"]).first()
        if not existing:
            new_type = QuestionType(
                name=type_data["name"],
                description=type_data["description"],
                code=type_data["code"],
                order_index=type_data["order_index"]
            )
            db.add(new_type)
    
    db.commit()

app = FastAPI(title="Mülakat Soru Hazırlama API", version="1.0.0")

# Startup event'te default soru tiplerini oluştur
@app.on_event("startup")
async def startup_event():
    db = next(get_db())
    create_default_question_types(db)
    db.close()

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Health check endpoint
@app.get("/health")
async def health_check():
    return {"status": "healthy", "service": "mulakat-backend"}

# Wizard Adım 1: İlan bilgilerini getir
@app.get("/api/step1/contract/{contract_id}")
async def get_contract(contract_id: int, db: Session = Depends(get_db)):
    """İlan bilgilerini getir"""
    from .models import Contract
    
    try:
        contract = db.query(Contract).filter(Contract.id == contract_id).first()
        
        if not contract:
            raise HTTPException(status_code=404, detail="İlan bulunamadı")
        
        return {
            "success": True,
            "contract": {
                "id": contract.id,
                "title": contract.title,
                "content": contract.content,
                "general_requirements": contract.general_requirements,
                "created_at": contract.created_at.isoformat()
            }
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# Wizard Adım 1: İlan kaydet
@app.post("/api/step1/save-contract")
async def save_contract(
    contract_data: Dict[str, Any],
    db: Session = Depends(get_db)
):
    """İlan bilgilerini kaydet"""
    from .models import Contract
    
    try:
        title = contract_data.get("title", "").strip()
        content = contract_data.get("content", "").strip()
        general_requirements = contract_data.get("general_requirements", "").strip()
        
        # Aynı ilan adı kontrolü
        existing_contract = db.query(Contract).filter(
            Contract.title == title
        ).first()
        
        if existing_contract:
            # Aynı isimle kayıt yapmayı engelle
            return {
                "success": False,
                "warning": True,
                "message": f"'{title}' adında bir ilan zaten mevcut. Lütfen farklı bir ilan adı girin.",
                "existing_contract": {
                    "id": existing_contract.id,
                    "title": existing_contract.title,
                    "created_at": existing_contract.created_at.isoformat()
                }
            }
        
        # Yeni ilan oluştur
        new_contract = Contract(
            title=title,
            content=content,
            general_requirements=general_requirements
        )
        
        db.add(new_contract)
        db.commit()
        db.refresh(new_contract)
        
        return {
            "success": True,
            "contract": {
                "id": new_contract.id,
                "title": new_contract.title,
                "content": new_contract.content,
                "general_requirements": new_contract.general_requirements
            },
            "message": "İlan başarıyla kaydedildi"
        }
        
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))

# Wizard Adım 2: Rolleri listele
@app.get("/api/step2/roles/{contract_id}")
async def get_roles(contract_id: int, db: Session = Depends(get_db)):
    """Belirli bir ilanın rollerini getir"""
    from .models import Role
    
    roles = db.query(Role).filter(Role.contract_id == contract_id).all()
    
    return {
        "success": True,
        "roles": [
            {
                "id": role.id,
                "name": role.name,
                "salary_multiplier": role.salary_multiplier,
                "position_count": role.position_count,
                "special_requirements": role.requirements
            } for role in roles
        ]
    }

# Wizard Adım 2: Yeni rol ekle
@app.post("/api/step2/add-role")
async def add_role(
    role_data: Dict[str, Any],
    db: Session = Depends(get_db)
):
    """Yeni rol ekle"""
    from .models import Role
    
    try:
        new_role = Role(
            contract_id=role_data.get("contract_id"),
            name=role_data.get("name"),
            salary_multiplier=role_data.get("salary_multiplier"),
            position_count=role_data.get("position_count"),
            requirements=role_data.get("special_requirements", "")
        )
        
        db.add(new_role)
        db.commit()
        db.refresh(new_role)
        
        return {
            "success": True,
            "role": {
                "id": new_role.id,
                "name": new_role.name,
                "salary_multiplier": new_role.salary_multiplier,
                "position_count": new_role.position_count,
                "special_requirements": new_role.requirements
            },
            "message": "Rol başarıyla eklendi"
        }
        
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))

# Wizard Adım 2: Rol güncelle
@app.put("/api/step2/roles/{role_id}")
async def update_role(
    role_id: int, 
    role_data: Dict[str, Any], 
    db: Session = Depends(get_db)
):
    """Rolü güncelle"""
    from .models import Role
    
    try:
        role = db.query(Role).filter(Role.id == role_id).first()
        
        if not role:
            raise HTTPException(status_code=404, detail="Rol bulunamadı")
        
        # Rol bilgilerini güncelle
        role.name = role_data.get("name", role.name)
        role.salary_multiplier = role_data.get("salary_multiplier", role.salary_multiplier)
        role.position_count = role_data.get("position_count", role.position_count)
        role.requirements = role_data.get("special_requirements", role.requirements)
        
        db.commit()
        db.refresh(role)
        
        return {
            "success": True,
            "role": {
                "id": role.id,
                "name": role.name,
                "salary_multiplier": role.salary_multiplier,
                "position_count": role.position_count,
                "special_requirements": role.requirements
            },
            "message": "Rol başarıyla güncellendi"
        }
        
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))

# Wizard Adım 2: Rol sil
@app.delete("/api/step2/roles/{role_id}")
async def delete_role(role_id: int, db: Session = Depends(get_db)):
    """Rol sil"""
    from .models import Role
    
    try:
        role = db.query(Role).filter(Role.id == role_id).first()
        
        if not role:
            raise HTTPException(status_code=404, detail="Rol bulunamadı")
        
        db.delete(role)
        db.commit()
        
        return {
            "success": True,
            "message": "Rol başarıyla silindi"
        }
        
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))

# Wizard Adım 3: Global sınav konfigürasyonu
@app.get("/api/step3/global-config/{contract_id}")
async def get_global_question_config(contract_id: int, db: Session = Depends(get_db)):
    """Global sınav ayarlarını getir"""
    from .models import QuestionConfig
    
    try:
        # Mevcut konfigürasyonu kontrol et
        config = db.query(QuestionConfig).filter(
            QuestionConfig.contract_id == contract_id
        ).first()
        
        if not config:
            # Default konfigürasyon oluştur
            config = QuestionConfig(
                contract_id=contract_id,
                candidate_multiplier=10,
                questions_per_candidate=5,
                question_type_distribution={
                    "professional_experience": 1,
                    "theoretical_knowledge": 2,
                    "practical_application": 2
                }
            )
            db.add(config)
            db.commit()
            db.refresh(config)
        
        # Aktif soru tiplerini de getir
        question_types = db.query(QuestionType).filter(
            QuestionType.is_active == True
        ).order_by(QuestionType.order_index).all()
        
        return {
            "success": True,
            "global_config": {
                "candidate_multiplier": config.candidate_multiplier,
                "questions_per_candidate": config.questions_per_candidate,
                "question_type_distribution": config.question_type_distribution or {}
            },
            "available_question_types": [{
                "id": qt.id,
                "name": qt.name,
                "code": qt.code,
                "description": qt.description
            } for qt in question_types]
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/step3/save-global-config")
async def save_global_question_config(
    config_data: Dict[str, Any],
    db: Session = Depends(get_db)
):
    """Global sınav ayarlarını kaydet ve mevcut rol konfigürasyonlarını yeniden hesapla"""
    from .models import QuestionConfig, RoleQuestionConfig, Role
    
    try:
        contract_id = config_data.get("contract_id")
        
        # Mevcut konfigürasyonu al veya oluştur
        config = db.query(QuestionConfig).filter(
            QuestionConfig.contract_id == contract_id
        ).first()
        
        if config:
            # Güncelle
            config.candidate_multiplier = config_data.get("candidate_multiplier", 10)
            config.questions_per_candidate = config_data.get("questions_per_candidate", 5)
            config.question_type_distribution = config_data.get("question_type_distribution", {})
        else:
            # Yeni oluştur
            config = QuestionConfig(
                contract_id=contract_id,
                candidate_multiplier=config_data.get("candidate_multiplier", 10),
                questions_per_candidate=config_data.get("questions_per_candidate", 5),
                question_type_distribution=config_data.get("question_type_distribution", {})
            )
            db.add(config)
        
        # ✅ KIRITIK: Global ayarlar değişti, tüm mevcut rol konfigürasyonlarını sil!
        # Bu sayede yeni hesaplama kullanılacak
        roles = db.query(Role).filter(Role.contract_id == contract_id).all()
        role_ids = [role.id for role in roles]
        
        if role_ids:
            # Bu contract'a ait tüm rol konfigürasyonlarını sil
            db.query(RoleQuestionConfig).filter(
                RoleQuestionConfig.role_id.in_(role_ids)
            ).delete(synchronize_session=False)
        
        db.commit()
        db.refresh(config)
        
        return {
            "success": True,
            "message": "Global sınav ayarları kaydedildi ve roller yeniden hesaplandı!",
            "global_config": {
                "candidate_multiplier": config.candidate_multiplier,
                "questions_per_candidate": config.questions_per_candidate,
                "question_type_distribution": config.question_type_distribution or {}
            },
            "reset_info": f"{len(role_ids)} rol konfigürasyonu sıfırlandı"
        }
        
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))

# Wizard Adım 3: Rollere göre soru konfigürasyonu getir
@app.get("/api/step3/role-question-configs/{contract_id}")
async def get_role_question_configs(contract_id: int, db: Session = Depends(get_db)):
    """Tüm rollerin soru konfigürasyonlarını getir (yeni hesaplama mantığı ile)"""
    from .models import Role, RoleQuestionConfig, QuestionConfig
    
    try:
        # Global sınav ayarlarını al
        global_config = db.query(QuestionConfig).filter(
            QuestionConfig.contract_id == contract_id
        ).first()
        
        if not global_config:
            # Default global config oluştur
            global_config = QuestionConfig(
                contract_id=contract_id,
                candidate_multiplier=10,
                questions_per_candidate=5,
                question_type_distribution={
                    "professional_experience": 1,
                    "theoretical_knowledge": 2,
                    "practical_application": 2
                }
            )
            db.add(global_config)
            db.commit()
            db.refresh(global_config)
        
        # Rolleri al
        roles = db.query(Role).filter(Role.contract_id == contract_id).all()
        
        # Aktif soru tiplerini al
        question_types = db.query(QuestionType).filter(
            QuestionType.is_active == True
        ).order_by(QuestionType.order_index).all()
        
        role_configs = []
        for role in roles:
            # Role ait konfigürasyonları al
            configs = db.query(RoleQuestionConfig).filter(
                RoleQuestionConfig.role_id == role.id
            ).all()
            
            # Konfigürasyonları soru tipine göre eşleştir
            config_map = {config.question_type_id: config for config in configs}
            
            question_type_configs = []
            for qt in question_types:
                config = config_map.get(qt.id)
                
                # Yeni hesaplama mantığı: pozisyon × aday_çarpanı × aday_başına_soru
                if config:
                    # Mevcut konfigürasyon varsa onu kullan
                    default_count = config.question_count
                else:
                    # Global config'e göre hesapla
                    candidate_count = role.position_count * global_config.candidate_multiplier
                    
                    # Dinamik soru türü dağılımından çek
                    distribution = global_config.question_type_distribution or {}
                    questions_per_candidate = distribution.get(qt.code, 1)  # Default 1
                    default_count = candidate_count * questions_per_candidate
                
                question_type_configs.append({
                    "question_type_id": qt.id,
                    "question_type_name": qt.name,
                    "question_type_description": qt.description,
                    "question_type_code": qt.code,
                    "question_count": default_count,
                    "difficulty_level": config.difficulty_level if config else "Orta"
                })
            
            # Hesaplama bilgileri
            candidate_count = role.position_count * global_config.candidate_multiplier
            
            role_data = {
                "role_id": role.id,
                "role_name": role.name,
                "salary_multiplier": role.salary_multiplier,
                "position_count": role.position_count,
                "candidate_count": candidate_count,  # Hesaplanan aday sayısı
                "question_types": question_type_configs
            }
            role_configs.append(role_data)
        
        return {
            "success": True,
            "role_configs": role_configs
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# Wizard Adım 3: Role soru konfigürasyonu kaydet
@app.post("/api/step3/save-role-question-config")
async def save_role_question_config(
    config_data: Dict[str, Any],
    db: Session = Depends(get_db)
):
    """Bir role ait soru konfigürasyonunu kaydet"""
    from .models import RoleQuestionConfig
    
    try:
        role_id = config_data.get("role_id")
        question_type_id = config_data.get("question_type_id")
        
        # Mevcut konfigürasyonu kontrol et
        existing_config = db.query(RoleQuestionConfig).filter(
            RoleQuestionConfig.role_id == role_id,
            RoleQuestionConfig.question_type_id == question_type_id
        ).first()
        
        if existing_config:
            # Güncelle
            existing_config.question_count = config_data.get("question_count", 5)
            existing_config.difficulty_level = "Orta"  # Maaş katsayısına göre belirlenecek
            
            db.commit()
            db.refresh(existing_config)
            config = existing_config
        else:
            # Yeni oluştur
            new_config = RoleQuestionConfig(
                role_id=role_id,
                question_type_id=question_type_id,
                question_count=config_data.get("question_count", 5),
                difficulty_level="Orta"  # Maaş katsayısına göre belirlenecek
            )
            
            db.add(new_config)
            db.commit()
            db.refresh(new_config)
            config = new_config
        
        return {
            "success": True,
            "config": {
                "id": config.id,
                "role_id": config.role_id,
                "question_type_id": config.question_type_id,
                "question_count": config.question_count,
                "difficulty_level": config.difficulty_level
            },
            "message": "Rol soru konfigürasyonu başarıyla kaydedildi"
        }
        
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))

# Wizard Adım 3: Tüm rol konfigürasyonlarını toplu kaydet
@app.post("/api/step3/save-all-role-configs")
async def save_all_role_configs(
    configs_data: Dict[str, Any],
    db: Session = Depends(get_db)
):
    """Tüm rollerin soru konfigürasyonlarını toplu kaydet"""
    from .models import RoleQuestionConfig
    
    try:
        contract_id = configs_data.get("contract_id")
        role_configs = configs_data.get("role_configs", [])
        
        saved_configs = []
        
        for role_config in role_configs:
            role_id = role_config.get("role_id")
            question_types = role_config.get("question_types", [])
            
            for qt_config in question_types:
                question_type_id = qt_config.get("question_type_id")
                question_count = qt_config.get("question_count", 5)
                # Zorluk seviyesi maaş katsayısına göre belirlenecek, şimdilik "Orta" default
                difficulty_level = "Orta"
                
                # Mevcut konfigürasyonu kontrol et
                existing_config = db.query(RoleQuestionConfig).filter(
                    RoleQuestionConfig.role_id == role_id,
                    RoleQuestionConfig.question_type_id == question_type_id
                ).first()
                
                if existing_config:
                    # Güncelle
                    existing_config.question_count = question_count
                    existing_config.difficulty_level = difficulty_level
                    
                    db.commit()
                    db.refresh(existing_config)
                    saved_config = existing_config
                else:
                    # Yeni oluştur
                    new_config = RoleQuestionConfig(
                        role_id=role_id,
                        question_type_id=question_type_id,
                        question_count=question_count,
                        difficulty_level=difficulty_level
                    )
                    
                    db.add(new_config)
                    db.commit()
                    db.refresh(new_config)
                    saved_config = new_config
                
                saved_configs.append({
                    "role_id": saved_config.role_id,
                    "question_type_id": saved_config.question_type_id,
                    "question_count": saved_config.question_count,
                    "difficulty_level": saved_config.difficulty_level
                })
        
        return {
            "success": True,
            "saved_configs": saved_configs,
            "message": f"{len(saved_configs)} soru konfigürasyonu başarıyla kaydedildi"
        }
        
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))

# Wizard Adım 4: Direkt soru üretimi (JSON adımı kaldırıldı)
@app.post("/api/step4/generate-questions")
async def generate_questions_directly(
    request_data: Dict[str, Any],
    db: Session = Depends(get_db)
):
    """Genel şartlar, özel şartlar ve konfigürasyona göre direkt soru üret"""
    from .utils import generate_questions_with_4o_mini
    from .models import Contract, Role, RoleQuestionConfig, QuestionType, Question
    
    try:
        contract_id = request_data.get("contract_id")
        model_name = request_data.get("model_name", "gpt-4o-mini")
        
        # Contract ve rolleri al
        contract = db.query(Contract).filter(Contract.id == contract_id).first()
        if not contract:
            raise HTTPException(status_code=404, detail="İlan bulunamadı")
        
        roles = db.query(Role).filter(Role.contract_id == contract_id).all()
        
        all_questions = []
        
        # Her rol için sorular üret
        for role in roles:
            # Rol konfigürasyonlarını al
            configs = db.query(RoleQuestionConfig).filter(
                RoleQuestionConfig.role_id == role.id
            ).all()
            
            question_types = db.query(QuestionType).filter(
                QuestionType.is_active == True
            ).order_by(QuestionType.order_index).all()
            
            config_map = {config.question_type_id: config for config in configs}
            
            # Debug: Hesaplama bilgilerini logla
            logger.info(f"Role: {role.name}")
            logger.info(f"Position count: {role.position_count}")
            logger.info(f"Configs found: {len(configs)}")
            logger.info(f"Config map keys: {list(config_map.keys())}")
            
            # Soru dağılımını hesapla (Step 3'teki mantıkla aynı)
            question_distribution = {}
            
            # Global config'i al
            global_config = db.query(QuestionConfig).filter(
                QuestionConfig.contract_id == contract_id
            ).first()
            
            if global_config:
                for qt in question_types:
                    config = config_map.get(qt.id)
                    
                    if config and hasattr(config, 'question_count'):
                        # Mevcut konfigürasyon varsa onu kullan
                        count = config.question_count
                    else:
                        # Global config'e göre hesapla (Step 3'teki mantık)
                        candidate_count = role.position_count * global_config.candidate_multiplier
                        distribution = global_config.question_type_distribution or {}
                        
                        # distribution bir dict değilse default değer kullan
                        if isinstance(distribution, dict):
                            questions_per_candidate = distribution.get(qt.code, 1)
                        else:
                            questions_per_candidate = 1
                            
                        count = candidate_count * questions_per_candidate
                    
                    question_distribution[qt.code] = count
                    logger.info(f"Question type {qt.code}: {count} questions")
            else:
                # Global config yoksa default değerler
                for qt in question_types:
                    config = config_map.get(qt.id)
                    count = config.question_count if (config and hasattr(config, 'question_count')) else 5
                    question_distribution[qt.code] = count
                    logger.info(f"Question type {qt.code}: {count} questions (default)")
            
            # Zorluk seviyesi hesapla
            role_difficulty = get_difficulty_level_by_multiplier(role.salary_multiplier)
            
            # Soru üretimi için context hazırla
            job_context = f"""
İLAN BAŞLIĞI: {contract.title}

GENEL ŞARTLAR:
{contract.general_requirements or "Genel şartlar belirtilmemiş"}

ROL: {role.name}
MAAŞ KATSAYISI: {role.salary_multiplier}x
POZİSYON SAYISI: {role.position_count}
ÖZEL ŞARTLAR:
{role.requirements or "Özel şartlar belirtilmemiş"}

ZORLUK SEVİYESİ: {role_difficulty['description']}
"""
            
            # 4o mini API ile sorular üret
            questions_result = generate_questions_with_4o_mini(
                model_name=model_name,
                job_context=job_context,
                roles=[{
                    "name": role.name,
                    "salary_multiplier": role.salary_multiplier,
                    "position_count": role.position_count,
                    "special_requirements": role.requirements
                }],
                question_config={
                    "professional_experience": question_distribution.get("professional_experience", 5),
                    "theoretical_knowledge": question_distribution.get("theoretical_knowledge", 5),
                    "practical_application": question_distribution.get("practical_application", 5),
                    "difficulty_level": role_difficulty["level"]
                }
            )
            
            if questions_result["success"]:
                # Soruları veritabanına kaydet
                questions = questions_result["questions"]
                for question_type, question_list in questions.items():
                    for q in question_list:
                        new_question = Question(
                            role_id=role.id,
                            contract_id=contract_id,
                            question_text=q["question"],
                            question_type=question_type,
                            difficulty=q["difficulty"],
                            expected_answer=q.get("expected_answer", ""),
                            scoring_criteria=q.get("scoring_criteria", ""),
                            llm_model=model_name
                        )
                        db.add(new_question)
                
                all_questions.append({
                    "role_name": role.name,
                    "role_id": role.id,
                    "salary_multiplier": role.salary_multiplier,
                    "questions": questions,
                    "difficulty_info": role_difficulty,
                    "model_used": model_name,
                    "gpu_used": questions_result.get("gpu_used", False)
                })
            else:
                # Hata durumunda
                all_questions.append({
                    "role_name": role.name,
                    "role_id": role.id,
                    "salary_multiplier": role.salary_multiplier,
                    "error": questions_result.get("error", "Soru üretiminde hata"),
                    "model_used": model_name,
                    "gpu_used": questions_result.get("gpu_used", False)
                })
        
        db.commit()
        
        return {
            "success": True,
            "questions": all_questions,
            "total_roles": len(roles),
            "model_used": model_name,
            "message": f"{len(roles)} rol için sorular üretildi."
        }
        
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))

# Soruları görüntüle
@app.get("/api/step4/questions/{contract_id}")
async def get_generated_questions(
    contract_id: int,
    db: Session = Depends(get_db)
):
    """Üretilen soruları getir"""
    from .models import Question, Role, Contract
    
    try:
        # Contract kontrolü
        contract = db.query(Contract).filter(Contract.id == contract_id).first()
        if not contract:
            raise HTTPException(status_code=404, detail="İlan bulunamadı")
        
        # Rolleri al
        roles = db.query(Role).filter(Role.contract_id == contract_id).all()
        
        questions_by_role = []
        
        for role in roles:
            # Bu role ait soruları al
            questions = db.query(Question).filter(
                Question.role_id == role.id,
                Question.contract_id == contract_id
            ).all()
            
            # Soruları tipine göre grupla
            questions_by_type = {
                "professional_experience": [],
                "theoretical_knowledge": [],
                "practical_application": []
            }
            
            for q in questions:
                questions_by_type[q.question_type].append({
                    "id": q.id,
                    "question": q.question_text,
                    "difficulty": q.difficulty,
                    "expected_answer": q.expected_answer,
                    "scoring_criteria": q.scoring_criteria,
                    "model_used": q.llm_model
                })
            
            questions_by_role.append({
                "role_name": role.name,
                "role_id": role.id,
                "questions": questions_by_type,
                "total_questions": len(questions)
            })
        
        return {
            "success": True,
            "contract_title": contract.title,
            "questions_by_role": questions_by_role,
            "total_roles": len(roles)
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# Soru Tipleri API
@app.get("/api/question-types")
async def get_question_types(db: Session = Depends(get_db)):
    """Aktif soru tiplerini getir"""
    try:
        question_types = db.query(QuestionType).filter(
            QuestionType.is_active == True
        ).order_by(QuestionType.order_index).all()
        
        return {
            "success": True,
            "question_types": [
                {
                    "id": qt.id,
                    "name": qt.name,
                    "description": qt.description,
                    "code": qt.code,
                    "order_index": qt.order_index
                }
                for qt in question_types
            ]
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/question-types")
async def create_question_type(
    question_type_data: Dict[str, Any],
    db: Session = Depends(get_db)
):
    """Yeni soru tipi oluştur"""
    try:
        # Aynı code'lu soru tipi var mı kontrol et
        existing = db.query(QuestionType).filter(
            QuestionType.code == question_type_data["code"]
        ).first()
        
        if existing:
            raise HTTPException(status_code=400, detail="Bu kod zaten kullanılıyor")
        
        new_question_type = QuestionType(
            name=question_type_data["name"],
            description=question_type_data.get("description", ""),
            code=question_type_data["code"],
            order_index=question_type_data.get("order_index", 0)
        )
        
        db.add(new_question_type)
        db.commit()
        db.refresh(new_question_type)
        
        return {
            "success": True,
            "question_type": {
                "id": new_question_type.id,
                "name": new_question_type.name,
                "description": new_question_type.description,
                "code": new_question_type.code,
                "order_index": new_question_type.order_index
            }
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.put("/api/question-types/{question_type_id}")
async def update_question_type(
    question_type_id: int,
    question_type_data: Dict[str, Any],
    db: Session = Depends(get_db)
):
    """Soru tipi güncelle"""
    try:
        question_type = db.query(QuestionType).filter(
            QuestionType.id == question_type_id
        ).first()
        
        if not question_type:
            raise HTTPException(status_code=404, detail="Soru tipi bulunamadı")
        
        # Code değiştiriliyorsa, aynı code'lu başka soru tipi var mı kontrol et
        if "code" in question_type_data and question_type_data["code"] != question_type.code:
            existing = db.query(QuestionType).filter(
                QuestionType.code == question_type_data["code"],
                QuestionType.id != question_type_id
            ).first()
            
            if existing:
                raise HTTPException(status_code=400, detail="Bu kod zaten kullanılıyor")
        
        # Güncelleme
        for key, value in question_type_data.items():
            setattr(question_type, key, value)
        
        db.commit()
        db.refresh(question_type)
        
        return {
            "success": True,
            "question_type": {
                "id": question_type.id,
                "name": question_type.name,
                "description": question_type.description,
                "code": question_type.code,
                "order_index": question_type.order_index
            }
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/api/question-types/{question_type_id}")
async def delete_question_type(
    question_type_id: int,
    db: Session = Depends(get_db)
):
    """Soru tipi sil (soft delete)"""
    try:
        question_type = db.query(QuestionType).filter(
            QuestionType.id == question_type_id
        ).first()
        
        if not question_type:
            raise HTTPException(status_code=404, detail="Soru tipi bulunamadı")
        
        # Kullanımda olan soru tipi mi kontrol et
        configs_count = db.query(RoleQuestionConfig).filter(
            RoleQuestionConfig.question_type_id == question_type_id
        ).count()
        
        if configs_count > 0:
            # Soft delete - sadece deaktif et
            question_type.is_active = False
            db.commit()
            return {"success": True, "message": "Soru tipi deaktif edildi"}
        else:
            # Hard delete - tamamen sil
            db.delete(question_type)
            db.commit()
            return {"success": True, "message": "Soru tipi silindi"}
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# Sistem Bilgileri API
@app.get("/api/system/4o-mini-models")
async def get_4o_mini_models():
    """Mevcut 4o mini modellerini getir"""
    from .utils import get_available_4o_mini_models
    
    try:
        models = get_available_4o_mini_models()
        return {
            "success": True,
            "models": models
        }
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "models": []
        }

# Eski Step 3 endpoint'lerini kaldır veya yorum yap
# Sistem bilgileri ve soru üretimi sonraki adımlarda kullanılacak

# Word dosyası oluşturma endpoint'i
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi.responses import FileResponse
import tempfile
import os

@app.post("/api/step5/generate-word")
async def generate_word_document(
    request_data: Dict[str, Any],
    db: Session = Depends(get_db)
):
    """Üretilen soruları ayrı Word dosyaları olarak ZIP içinde indir"""
    try:
        logger.info(f"Word dosyaları oluşturma başlatıldı. Request data: {request_data}")
        contract_id = request_data.get("contract_id")
        
        if not contract_id:
            raise HTTPException(status_code=400, detail="contract_id gerekli")
        
        logger.info(f"Contract ID: {contract_id}")
        
        # Contract bilgilerini al
        contract = db.query(Contract).filter(Contract.id == contract_id).first()
        if not contract:
            raise HTTPException(status_code=404, detail="İlan bulunamadı")
        
        logger.info(f"Contract bulundu: {contract.title}")
        
        # Rolleri al
        roles = db.query(Role).filter(Role.contract_id == contract_id).all()
        logger.info(f"Roller bulundu: {len(roles)} adet")
        
        # ZIP dosyası oluştur
        import zipfile
        import io
        
        zip_buffer = io.BytesIO()
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            # Her rol için dosyaları oluştur
            for role in roles:
                logger.info(f"Rol işleniyor: {role.name}")
                
                # Bu role ait soruları al
                questions = db.query(Question).filter(
                    Question.role_id == role.id,
                    Question.contract_id == contract_id
                ).all()
                logger.info(f"Rol {role.name} için {len(questions)} soru bulundu")
                
                # Soruları türlerine göre grupla
                questions_by_type = {}
                for q in questions:
                    if q.question_type not in questions_by_type:
                        questions_by_type[q.question_type] = []
                    questions_by_type[q.question_type].append(q)
                
                # Tür isimleri
                type_names = {
                    'professional_experience': 'Mesleki Deneyim Soruları',
                    'theoretical_knowledge': 'Teorik Bilgi Soruları',
                    'practical_application': 'Pratik Uygulama Soruları'
                }
                
                # Her aday için soru dosyası oluştur
                max_questions_per_type = max(len(q_list) for q_list in questions_by_type.values()) if questions_by_type else 0
                logger.info(f"Maksimum soru sayısı: {max_questions_per_type}")
                
                for candidate_num in range(1, max_questions_per_type + 1):
                    logger.info(f"Aday {candidate_num} için dosyalar oluşturuluyor...")
                    
                    # Soru dosyası (S) - Sadece sorular
                    doc_s = Document()
                    title_s = doc_s.add_heading('MÜLAKAT SORULARI', 0)
                    title_s.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Contract bilgileri
                    doc_s.add_heading('İlan Bilgileri', level=1)
                    doc_s.add_paragraph(f'İlan Adı: {contract.title}')
                    doc_s.add_paragraph(f'Oluşturulma Tarihi: {contract.created_at.strftime("%d.%m.%Y") if contract.created_at else "Belirtilmemiş"}')
                    doc_s.add_paragraph(f'Pozisyon: {role.name} ({role.salary_multiplier}x)')
                    doc_s.add_paragraph(f'Aday No: {candidate_num}')
                    doc_s.add_paragraph()
                    
                    # Bu aday için soruları ekle
                    role_title = f"{role.name} (Aylık brüt sözleşme ücret tavanının {role.salary_multiplier} katına kadar)"
                    doc_s.add_heading(role_title, level=2)
                    
                    for q_type, q_list in questions_by_type.items():
                        if len(q_list) >= candidate_num:
                            doc_s.add_heading(type_names.get(q_type, q_type), level=3)
                            question = q_list[candidate_num - 1]  # 0-indexed
                            p = doc_s.add_paragraph()
                            p.add_run(f'1. ').bold = True
                            p.add_run(question.question_text)
                            doc_s.add_paragraph()
                    
                    # Türkçe karakterleri temizle
                    safe_role_name = role.name.replace("Ş", "S").replace("Ç", "C").replace("Ğ", "G").replace("İ", "I").replace("Ö", "O").replace("Ü", "U").replace("ş", "s").replace("ç", "c").replace("ğ", "g").replace("ı", "i").replace("ö", "o").replace("ü", "u")
                    
                    # Soru dosyasını ZIP'e ekle
                    s_filename = f"{safe_role_name} {role.salary_multiplier}x S{candidate_num}.docx"
                    s_buffer = io.BytesIO()
                    doc_s.save(s_buffer)
                    s_buffer.seek(0)
                    zip_file.writestr(s_filename, s_buffer.getvalue())
                    logger.info(f"Soru dosyası eklendi: {s_filename}")
                    
                    # Cevap dosyası (C) - Sorular ve cevaplar
                    doc_c = Document()
                    title_c = doc_c.add_heading('MÜLAKAT SORULARI VE CEVAPLARI', 0)
                    title_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Contract bilgileri
                    doc_c.add_heading('İlan Bilgileri', level=1)
                    doc_c.add_paragraph(f'İlan Adı: {contract.title}')
                    doc_c.add_paragraph(f'Oluşturulma Tarihi: {contract.created_at.strftime("%d.%m.%Y") if contract.created_at else "Belirtilmemiş"}')
                    doc_c.add_paragraph(f'Pozisyon: {role.name} ({role.salary_multiplier}x)')
                    doc_c.add_paragraph(f'Aday No: {candidate_num}')
                    doc_c.add_paragraph()
                    
                    # Bu aday için soruları ve cevapları ekle
                    doc_c.add_heading(role_title, level=2)
                    
                    for q_type, q_list in questions_by_type.items():
                        if len(q_list) >= candidate_num:
                            doc_c.add_heading(type_names.get(q_type, q_type), level=3)
                            question = q_list[candidate_num - 1]  # 0-indexed
                            p = doc_c.add_paragraph()
                            p.add_run(f'1. ').bold = True
                            p.add_run(question.question_text)
                            if question.expected_answer:
                                answer_para = doc_c.add_paragraph()
                                answer_para.add_run('Beklenen Cevap: ').bold = True
                                answer_para.add_run(question.expected_answer)
                            doc_c.add_paragraph()
                    
                    # Cevap dosyasını ZIP'e ekle
                    c_filename = f"{safe_role_name} {role.salary_multiplier}x C{candidate_num}.docx"
                    c_buffer = io.BytesIO()
                    doc_c.save(c_buffer)
                    c_buffer.seek(0)
                    zip_file.writestr(c_filename, c_buffer.getvalue())
                    logger.info(f"Cevap dosyası eklendi: {c_filename}")
        
        # ZIP buffer'ı hazırla
        zip_buffer.seek(0)
        
        logger.info("ZIP dosyası oluşturuldu, Response döndürülüyor...")
        
        # Türkçe karakterleri temizle
        safe_title = contract.title.replace(" ", "_").replace("Ş", "S").replace("Ç", "C").replace("Ğ", "G").replace("İ", "I").replace("Ö", "O").replace("Ü", "U").replace("ş", "s").replace("ç", "c").replace("ğ", "g").replace("ı", "i").replace("ö", "o").replace("ü", "u")
        
        # ZIP dosyasını döndür
        return Response(
            content=zip_buffer.getvalue(),
            media_type='application/zip',
            headers={
                'Content-Disposition': f'attachment; filename="mulakat_sorulari_{safe_title}_{contract_id}.zip"'
            }
        )
        
    except Exception as e:
        logger.error(f"Word dosyaları oluşturma hatası: {str(e)}")
        logger.error(f"Hata detayı: {type(e).__name__}")
        import traceback
        logger.error(f"Traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    uvicorn.run("app.main:app", host="0.0.0.0", port=8000, reload=True) 